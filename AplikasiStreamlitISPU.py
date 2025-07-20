import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.gridspec import SubplotSpec 
from sklearn.preprocessing import MinMaxScaler
from scipy.stats import norm, weibull_min
import seaborn as sns
import warnings
import os
import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

warnings.filterwarnings('ignore')

st.set_page_config(page_title="Simulasi ISPU", layout="wide", page_icon="🌫️")
# ====== Sidebar Navigation ======
st.sidebar.title("Navigasi Halaman")
page = st.sidebar.radio("Pilih Halaman:", 
                       ["Simulasi ISPU", "Analisis Data", "Visualisasi Lanjutan"])

# ====== Main Page Content ======
if page == "Simulasi ISPU":
    

    # Title
    st.title("Laporan Simulasi ISPU dan Visualisasi")

    # ====== Fungsi untuk Dokumen Word ======
    def save_to_word(document, title, content=None, image=None, table=None):
        document.add_heading(title, level=2)
        if content:
            for line in content.split('\n'):
                if line.strip():
                    p = document.add_paragraph(line.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if table is not None:
            df = table
            t = document.add_table(rows=df.shape[0]+1, cols=df.shape[1])
            t.style = 'Table Grid'
            for j in range(df.shape[1]):
                t.cell(0, j).text = df.columns[j]
            for i in range(df.shape[0]):
                for j in range(df.shape[1]):
                    t.cell(i+1, j).text = str(df.values[i, j])
        if image:
            image_stream = BytesIO()
            plt.savefig(image_stream, format='png')
            plt.close()
            document.add_picture(image_stream, width=Inches(5))
        document.add_paragraph()

    def create_word_document(df, params, n_samples):
        doc = Document()
        doc.add_heading("Laporan Simulasi ISPU dan Visualisasi", 0)
        
        # ====== BAGIAN BARU: STATISTIK & PARAMETER DISTRIBUSI ======
        doc.add_heading("Statistik dan Parameter Distribusi", level=2)
        
        # 1. Statistik Deskriptif
        p = doc.add_paragraph()
        p.add_run("Statistik Deskriptif Simulasi (µg/m³)").bold = True
        p.add_run("\nBerikut ringkasan statistik data simulasi:")
        
        # Buat tabel statistik
        stats_table = df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].describe().round(2)
        t = doc.add_table(stats_table.shape[0]+1, stats_table.shape[1]+1)
        t.style = 'Table Grid'
        
        # Header kolom
        for j in range(stats_table.shape[1]):
            t.cell(0,j+1).text = stats_table.columns[j]
        
        # Header baris dan data
        for i in range(stats_table.shape[0]):
            t.cell(i+1,0).text = stats_table.index[i]
            for j in range(stats_table.shape[1]):
                t.cell(i+1,j+1).text = str(stats_table.iloc[i,j])
        
        doc.add_paragraph("Catatan: CO ditampilkan dalam µg/m³ (tanpa konversi)")
        
        # 2. Parameter Distribusi
        doc.add_paragraph().add_run("\nParameter Distribusi (Dihitung dari Data Historis)").bold = True

        # Buat tabel parameter dengan format yang berbeda untuk Normal dan Weibull
        param_table = doc.add_table(rows=7, cols=4)  # 4 kolom: Polutan, Distribusi, Parameter 1, Parameter 2
        param_table.style = 'Table Grid'

        # Header
        headers = ["Polutan", "Distribusi", "Parameter 1", "Parameter 2"]
        for j, header in enumerate(headers):
            param_table.cell(0,j).text = header

        # Isi data
        for i, col in enumerate(['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3'], start=1):
            dist_type, *dist_params = st.session_state.params_simulasi[col]
            param_table.cell(i,0).text = col
            param_table.cell(i,1).text = dist_type
            
            if dist_type == 'Weibull':
                c, loc, scale = dist_params
                param_table.cell(i,2).text = f"Shape (k): {c:.2f}"
                param_table.cell(i,3).text = f"Scale (λ): {scale:.2f}"
            else:  # Normal
                mu, std = dist_params
                param_table.cell(i,2).text = f"Mean: {mu:.2f}"
                param_table.cell(i,3).text = f"Std: {std:.2f}"

        doc.add_paragraph("Catatan:")
        doc.add_paragraph("- Weibull: Shape (k) dan Scale (λ)")
        doc.add_paragraph("- Normal: Mean dan Standard Deviation")
        
        # Distribusi Parameter
        plt.figure(figsize=(15, 10))
        plt.suptitle('Distribusi Parameter Kualitas Udara', fontsize=14)
        for i, col in enumerate(['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3'], 1):
            if col in params:
                dist_type, *dist_params = params[col]
                plt.subplot(2, 3, i)
                sns.histplot(df[col], bins=20, kde=False, stat="density", alpha=0.6, label='Histogram')
                
                x = np.linspace(0, df[col].max(), 100)
                if dist_type == 'Normal':
                    mu, std = dist_params
                    pdf = norm.pdf(x, mu, std)
                    plt.plot(x, pdf, 'g-', lw=2, label='Normal Fit')
                elif dist_type == 'Weibull':
                    c, loc, scale = dist_params
                    pdf = weibull_min.pdf(x, c, loc, scale)
                    plt.plot(x, pdf, 'r-', lw=2, label='Weibull Fit')
                    
                plt.title(f'{col} ({dist_type})')
                plt.xlabel('Konsentrasi')
                plt.ylabel('Density')
                plt.legend()
        plt.tight_layout(rect=[0, 0, 1, 0.95])
        save_to_word(doc, "Distribusi Parameter dengan Kurva Fit", 
                    content="""**Histogram menunjukkan distribusi konsentrasi polutan. Garis menunjukkan kurva distribusi yang difitkan.**
                            **Penjelasan:**
                            - Grafik membandingkan distribusi aktual (histogram) dengan kurva teoritis
                            - Garis hijau: Distribusi Normal
                            - Garis merah: Distribusi Weibull
                            - PM2.5, PM10, dan O3 dimodelkan dengan distribusi Weibull
                            - SO2, NO2, dan CO dimodelkan dengan distribusi normal
                            - Distribusi Weibull cocok untuk data dengan skewness positif
                    """,
                    image=True)
        
        # Tren Harian ISPU
        plt.figure(figsize=(14, 6))
        plt.plot(df.index, df['ISPU_max'], label='ISPU Harian', alpha=0.5)
        plt.plot(df['ISPU_max'].rolling(7).mean(), label='Rata-rata 7 Hari', color='red')
        plt.title('Tren Harian ISPU dengan Smoothing')
        plt.xlabel('Hari')
        plt.ylabel('ISPU')
        plt.grid(True)
        plt.legend()
        save_to_word(doc, "Tren Harian ISPU", 
                    content="""**Plot menunjukkan tren ISPU harian selama simulasi dengan rata-rata 7 hari terakhir.**
                            **Penjelasan Tren ISPU:**
                            - Garis biru: Nilai ISPU harian dengan fluktuasi alami
                            - Garis merah: Rata-rata 7 hari untuk melihat tren
                            - Lonjakan ISPU bisa disebabkan oleh:
                            - Peningkatan aktivitas kendaraan
                            - Pembakaran biomassa
                            - Kondisi meteorologi yang tidak mendispersi polutan
                    """,
                    image=True)
        
        # Kategori ISPU
        df['Kategori_ISPU'] = df['ISPU_max'].apply(lambda x: "Baik" if x <= 50 else 
                                                "Sedang" if x <= 100 else 
                                                "Tidak Sehat" if x <= 200 else 
                                                "Sangat Tidak Sehat" if x <= 300 else 
                                                "Berbahaya")

        kategori_counts = df['Kategori_ISPU'].value_counts()
        warna_kategori = {
            'Baik': 'green',
            'Sedang': 'gold',
            'Tidak Sehat': 'orange',
            'Sangat Tidak Sehat': 'red',
            'Berbahaya': 'darkred'
        }

        # Filter out categories with 0 count
        filtered_counts = kategori_counts[kategori_counts > 0]
        filtered_colors = [warna_kategori[k] for k in filtered_counts.index]

        # Create visualizations only if there are categories to show
        if not filtered_counts.empty:
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))
            
            # Pie Chart
            ax1.pie(filtered_counts, labels=filtered_counts.index,
                    autopct='%1.1f%%', startangle=90, colors=filtered_colors,
                    shadow=True, textprops={'color': 'black', 'fontsize': 10})
            ax1.set_title('Distribusi Persentase Kategori ISPU (Tanpa Kategori 0%)')
            ax1.axis('equal')
            
            # Bar Chart
            bars = ax2.bar(filtered_counts.index, filtered_counts.values, color=filtered_colors, edgecolor='black')
            ax2.set_title('Jumlah Hari per Kategori ISPU (Tanpa Kategori 0%)')
            ax2.set_xlabel('Kategori ISPU')
            ax2.set_ylabel('Jumlah Hari')
            ax2.grid(axis='y', linestyle='--', alpha=0.7)
            ax2.set_axisbelow(True)
            for bar in bars:
                height = bar.get_height()
                ax2.text(bar.get_x() + bar.get_width()/2., height + 0.5, 
                        f'{int(height)}\n({height/len(df):.1%})',
                        ha='center', va='bottom', fontsize=9)
            
            plt.tight_layout()
            
            # Save to Word document
            save_to_word(doc, "Distribusi Kategori ISPU", 
                        content="""**Diagram lingkaran dan batang menunjukkan distribusi kategori ISPU selama periode simulasi (kategori dengan 0% tidak ditampilkan).**
                                **Distribusi Kategori Kualitas Udara:**
                                - **Baik (0-50)**: Tidak ada risiko kesehatan
                                - **Sedang (51-100)**: Kelompok sensitif mungkin terpengaruh
                                - **Tidak Sehat (101-200)**: Seluruh populasi mulai terpengaruh
                                - **Sangat Tidak Sehat (201-300)**: Peringatan kesehatan serius
                                - **Berbahaya (>300)**: Darurat kesehatan publik
                        """,
                        image=True)
        else:
            st.warning("Tidak ada data kategori ISPU yang tersedia untuk ditampilkan.")

            # Polutan Dominan
            polutan_dominan = df['Parameter_Dominan'].value_counts()
            plt.figure(figsize=(10, 6))
            polutan_dominan.plot(kind='bar', color='skyblue')
            plt.title('Frekuensi Polutan Dominan Penyebab ISPU Tertinggi')
            plt.xlabel('Polutan')
            plt.ylabel('Jumlah Hari')
            plt.grid(axis='y', linestyle='--', alpha=0.7)
            save_to_word(doc, "Polutan Dominan Penyebab ISPU Tertinggi",
                        content="""**Diagram batang menunjukkan polutan mana yang paling sering menjadi penyebab ISPU tertinggi.**
                                **Penjelasan Diagram:**
                                - **Sumbu X**: Nama polutan (PM2.5, PM10, SO2, NO2, CO, O3)
                                - **Sumbu Y**: Jumlah hari polutan tersebut menjadi penyebab ISPU tertinggi
                                - Warna biru muda: Frekuensi dominansi masing-masing polutan
                                - Garis grid: Membantu membaca nilai secara akurat
                        """,
                        image=True)
            
            # Line Chart Polutan Utama
            plt.figure(figsize=(14, 6))
            for polutan in ['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']:
                plt.plot(df.index, df[polutan], label=polutan, alpha=0.7)
            plt.title('Tren Polutan Utama')
            plt.xlabel('Hari')
            plt.ylabel('Konsentrasi')
            plt.legend()
            plt.grid(True)
            save_to_word(doc, "Tren Polutan Utama",
                        content="""**Line chart menunjukkan tren beberapa polutan utama selama periode simulasi.**
                                **Penjelasan Diagram:**
                                - **Sumbu X**: Hari ke- (dikonversi dari hari simulasi)
                                - **Sumbu Y**: Konsentrasi harian (µg/m³)
                                - Garis berwarna: Mewakili polutan berbeda
                                - Area terarsir: Variasi harian dalam Hari tersebut
                                - Pola mingguan membantu identifikasi pengaruh aktivitas manusia
                        """,
                        image=True)

        #Histogram PM2.5
        plt.figure(figsize=(10, 6))
        sns.histplot(df['PM2.5'], bins=30, kde=True, color='teal')
        plt.axvline(x=15.5, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi PM2.5')
        plt.xlabel('PM2.5 (µg/m³)')
        plt.ylabel('Frekuensi (Jumlah Hari)')
        plt.legend()
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        save_to_word(doc, "Distribusi Konsentrasi PM2.5",
                    content="""**Histogram menunjukkan Distribusi PM2.5.**
                            **Penjelasan Diagram:**
                            - **Sumbu X**: Rentang konsentrasi PM2.5 (µg/m³)
                            - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
                            - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
                            - Garis KDE (Kernel Density Estimate): Estimasi kurva distribusi
                            - Garis Merah batas aman ISPU
                    """,
                    image=True)

        #Histogram PM10 
        plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['PM10'], bins=30, kde=True, color='teal')
        plt.axvline(x=50, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi PM10')
        plt.xlabel('PM10 (µg/m³)')
        plt.ylabel('Frekuensi (Jumlah Hari)')
        plt.legend()
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        save_to_word(doc, "Distribusi Konsentrasi PM10",
                    content="""**Histogram menunjukkan Distribusi PM10.**
                            **Penjelasan Diagram:**
                            - **Sumbu X**: Rentang konsentrasi PM10 (µg/m³)
                            - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
                            - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
                            - Garis KDE (Kernel Density Estimate): Estimasi kurva distribusi
                            - Garis Merah batas aman ISPU
                    """,
                    image=True)

        #Histogram SO2
        plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['SO2'], bins=30, kde=True, color='teal')
        plt.axvline(x=52, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi SO2')
        plt.xlabel('SO2 (µg/m³)')
        plt.ylabel('Frekuensi (Jumlah Hari)')
        plt.legend()
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        save_to_word(doc, "Distribusi Konsentrasi SO2",
                    content="""**Histogram menunjukkan Distribusi SO2.**
                            **Penjelasan Diagram:**
                            - **Sumbu X**: Rentang konsentrasi SO2 (µg/m³)
                            - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
                            - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
                            - Garis KDE: Estimasi kurva distribusi
                            - Garis Merah batas aman ISPU
                    """,
                    image=True)

        #Histogram NO2
        plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['NO2'], bins=30, kde=True, color='teal')
        plt.axvline(x=80, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi NO2')
        plt.xlabel('NO2 (µg/m³)')
        plt.ylabel('Frekuensi (Jumlah Hari)')
        plt.legend()
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        save_to_word(doc, "Distribusi Konsentrasi NO2",
                    content="""**Histogram menunjukkan Distribusi NO2.**
                            **Penjelasan Diagram:**
                            - **Sumbu X**: Rentang konsentrasi NO2 (µg/m³)
                            - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
                            - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
                            - Garis KDE: Estimasi kurva distribusi
                            - Garis Merah batas aman ISPU
                    """,
                    image=True)

        #Histogram CO
        plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['CO'], bins=30, kde=True, color='teal')
        plt.axvline(x=4, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi CO')
        plt.xlabel('CO (µg/m³)')
        plt.ylabel('Frekuensi (Jumlah Hari)')
        plt.legend()
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        save_to_word(doc, "Distribusi Konsentrasi CO",
                    content="""**Histogram menunjukkan Distribusi CO.**
                            **Penjelasan Diagram:**
                            - **Sumbu X**: Rentang konsentrasi CO (mg/m³)
                            - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
                            - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
                            - Garis KDE: Estimasi kurva distribusi
                            - Garis Merah batas aman ISPU
                    """,
                    image=True)

        #Histogram O3
        plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['O3'], bins=30, kde=True, color='teal')
        plt.axvline(x=120, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi O3')
        plt.xlabel('O3 (µg/m³)')
        plt.ylabel('Frekuensi (Jumlah Hari)')
        plt.legend()
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        save_to_word(doc, "Distribusi Konsentrasi O3",
                    content="""**Histogram menunjukkan Distribusi O3.**
                            **Penjelasan Diagram:**
                            - **Sumbu X**: Rentang konsentrasi O3 (µg/m³)
                            - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
                            - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
                            - Garis KDE: Estimasi kurva distribusi
                            - Garis Merah batas aman ISPU
                    """,
                    image=True)

        # Kesimpulan
        rata2_ispu = df['ISPU_max'].mean()
        max_ispu = df['ISPU_max'].max()
        hari_tidak_sehat = len(df[df['Kategori_ISPU'].isin(['Tidak Sehat', 'Sangat Tidak Sehat', 'Berbahaya'])])
        polusi_tertinggi = df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].max().idxmax()
        hari_dengan_polusi_tinggi = len(df[df['Polusi_Total'] > df['Polusi_Total'].quantile(0.75)])


        pm25_over_limit = len(st.session_state.df[st.session_state.df['PM2.5'] > 15.5])
        pm10_over_limit = len(st.session_state.df[st.session_state.df['PM10'] > 50])
        o3_over_limit = len(st.session_state.df[st.session_state.df['O3'] > 120])
        co_over_limit = len(st.session_state.df[st.session_state.df['CO'] > 5])
        so2_over_limit = len(st.session_state.df[st.session_state.df['SO2'] > 20])
        no2_over_limit = len(st.session_state.df[st.session_state.df['NO2'] > 30])
        max_pollutant = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].max().idxmax()
        min_pollutant = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].min().idxmax()
        kategori_urut = ['Baik', 'Sedang', 'Tidak Sehat', 'Sangat Tidak Sehat', 'Berbahaya']
        kategori_counts = st.session_state.df['Kategori_ISPU'].value_counts().reindex(kategori_urut, fill_value=0)
        mean_pollutants = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].mean().sort_values()
        least_harmful = mean_pollutants.idxmin()
        min_pollutants = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].min().sort_values()
        cleanest_pollutant = min_pollutants.idxmin()
        std_pollutants = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].std().sort_values()
        most_stable = std_pollutants.idxmin()

        kesimpulan = f"""
        1. Rata-rata nilai ISPU selama {st.session_state.n_samples} hari adalah {rata2_ispu:.2f}, menunjukkan bahwa kualitas udara cenderung berada pada kategori {'Baik/Sedang' if rata2_ispu < 100 else 'Tidak Sehat'}.
        2. Nilai ISPU tertinggi adalah {max_ispu:.2f}, masuk dalam kategori '{st.session_state.df.loc[st.session_state.df['ISPU_max'].idxmax(), 'Kategori_ISPU']}'. Ini merupakan kondisi kritis yang memerlukan antisipasi lebih lanjut.
        3. Terdapat {hari_tidak_sehat} hari dengan kategori ISPU tidak sehat/sangat tidak sehat/berbahaya. Kelompok rentan seperti lansia dan penderita penyakit paru-paru perlu waspada.
        4. Polutan dominan berdasarkan konsentrasi tertinggi adalah '{polusi_tertinggi}', menunjukkan perlunya fokus mitigasi pada parameter tersebut.
        5. Terdapat {hari_dengan_polusi_tinggi} hari dengan polusi total sangat tinggi (>75% kuartil), yang mungkin dipengaruhi oleh peningkatan beberapa polutan secara bersamaan.
        6. PM2.5 melebihi batas aman ISPU (15.5 µg/m³) sebanyak {pm25_over_limit} hari, menunjukkan perlunya pengendalian emisi kendaraan dan industri.
        7. PM10 melebihi batas aman ISPU (50 µg/m³) sebanyak {pm10_over_limit} hari, menunjukkan polusi udara yang cukup signifikan di wilayah simulasi.
        8. Ozon (O3) melebihi ambang batas aman sebanyak {o3_over_limit} hari.
        9. CO melebihi ambang batas sebanyak {co_over_limit} hari.
        10. SO2 melebihi ambang batas sebanyak {so2_over_limit}.
        11. NO2 melebihi ambang batas sebanyak {no2_over_limit}.
        12. Polutan {max_pollutant} memiliki konsentrasi tertinggi di seluruh simulasi, menunjukkan bahwa polutan tersebut paling berkontribusi terhadap penurunan kualitas udara.
        13. Polutan {min_pollutant} memiliki konsentrasi tertinggi di minimim.
        14. {least_harmful} merupakan polutan dengan konsentrasi rata-rata terendah (**{mean_pollutants[least_harmful]:.2f} µg/m³**). Sumber emisi {least_harmful} relatif lebih terkendali.  
        15. {cleanest_pollutant} mencapai level terendah (**{min_pollutants[cleanest_pollutant]:.2f} µg/m³**) saat udara paling bersih.
        16. {most_stable} memiliki fluktuasi harian paling stabil (deviasi **{std_pollutants[most_stable]:.2f} µg/m³**).  
        """
        
        rekomendasi = """
        **Rekomendasi:**
        - Pengurangan emisi kendaraan bermotor dan aktivitas industri
        - Edukasi publik tentang cara melindungi diri dari paparan polusi udara
        - Peningkatan pemantauan kualitas udara di area dengan risiko tinggi
        - Fokus pada pengendalian polutan dominan yang teridentifikasi
        - Implementasi sistem peringatan dini ketika konsentrasi polutan mencapai level berbahaya
        - Fokus pada pengendalian polutan dengan **rata-rata tertinggi** terlebih dahulu  
        - Manfaatkan data polutan **paling stabil** sebagai baseline pemantauan  
        - Gunakan kriteria **minimum terendah** untuk menetapkan target kualitas udara bersih 
        """
        
        save_to_word(doc, "Kesimpulan Simulasi", content=kesimpulan)
        save_to_word(doc, "Rekomendasi", content=rekomendasi)
        
        # Simpan ke buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer

    # ====== Inisialisasi Session State ======
    if 'df' not in st.session_state:
        st.session_state.df = None
    if 'params' not in st.session_state:
        st.session_state.params = None
    if 'df_excel' not in st.session_state:
        st.session_state.df_excel = None
    if 'n_samples' not in st.session_state:
        st.session_state.n_samples = 365

    # ====== STEP 1: BACA FILE EXCEL DAN FIT PARAMETER DISTRIBUSI ======
    @st.cache_data
    def load_data():
        file_excel = "TugasBesar.xlsx"  # Update with your file path
        df_excel = pd.read_excel(file_excel, sheet_name="Sort")
        
        # Filter hanya untuk baris dengan stasiun 'DKI1 Bunderan HI'
        df_excel = df_excel[df_excel['stasiun'] == 'DKI1 Bunderan HI']
        
        # Ambil dan ganti nama kolom sesuai format standar
        df_excel = df_excel[['pm_sepuluh', 'pm_duakomalima', 'sulfur_dioksida', 'karbon_monoksida', 'ozon', 'nitrogen_dioksida']].copy()
        df_excel.columns = ['PM10', 'PM2.5', 'SO2', 'CO', 'O3', 'NO2']
        
        # Konversi ke numerik dan hapus baris NaN
        df_excel = df_excel.apply(pd.to_numeric, errors='coerce').dropna()
        return df_excel

    if st.session_state.df_excel is None:
        st.session_state.df_excel = load_data()

    # Estimasi parameter distribusi untuk semua polutan
    if st.session_state.params is None:
        st.session_state.params = {}
        
        # Fit Weibull untuk PM2.5, PM10, dan O3
        for col in ['PM2.5', 'PM10', 'O3']:
            data = st.session_state.df_excel[col]
            params = weibull_min.fit(data, floc=0)  # Force location parameter to 0
            st.session_state.params[col] = ('Weibull', *params)
        
        # Fit Normal untuk SO2, NO2, dan CO
        for col in ['SO2', 'NO2', 'CO']:
            mu, std = norm.fit(st.session_state.df_excel[col])
            st.session_state.params[col] = ('Normal', mu, std)

    # ====== STEP 2: GENERATE DATA SIMULASI ======
    def generate_simulation(n_samples=365):
        rng = np.random.default_rng()
        df = pd.DataFrame()
        
        for col in st.session_state.df_excel.columns:
            dist_type, *dist_params = st.session_state.params[col]
            
            if dist_type == 'Weibull':
                # Generate data Weibull untuk PM2.5, PM10, dan O3
                c, loc, scale = dist_params
                data = weibull_min.rvs(c, loc, scale, size=n_samples, random_state=rng)
                
                # Handle nilai terlalu kecil dengan refleksi
                small_indices = data < 0.1  # Threshold untuk nilai sangat kecil
                if np.any(small_indices):
                    # Generate nilai baru dari distribusi yang sama
                    new_values = weibull_min.rvs(c, loc, scale, size=np.sum(small_indices), random_state=rng)
                    # Pastikan nilai baru tidak terlalu kecil
                    new_values = np.where(new_values < 0.1, new_values + 0.1, new_values)
                    data[small_indices] = new_values
                    
            else:  # Distribusi Normal untuk SO2, NO2, CO
                mu, std = dist_params
                # 1. Generate data normal
                data = rng.normal(mu, std, n_samples)
                
                # 2. Handle nilai negatif dengan lebih elegan
                neg_indices = data < 0
                
                if np.any(neg_indices):
                    # Hitung probabilitas kumulatif untuk nilai negatif
                    cdf_neg = norm.cdf(0, loc=mu, scale=std)
                    
                    # Generate nilai baru dari truncated normal (0 sampai infinity)
                    truncated_values = mu + std * rng.standard_normal(size=np.sum(neg_indices))
                    truncated_values = np.abs(truncated_values)  # Refleksi nilai
                    
                    # Gabungkan dengan data positif
                    data[neg_indices] = truncated_values
            
            df[col] = np.round(data, 2)
        
        return df

    # User input for number of samples
    st.session_state.n_samples = st.sidebar.number_input("Jumlah Sampel Hari", min_value=30, max_value=1000, value=365)

    if st.session_state.df is None:
        st.session_state.df = generate_simulation(st.session_state.n_samples)

    # Fungsi untuk menghitung ulang data jika n_samples berubah
    def update_simulation():
        st.session_state.df = generate_simulation(st.session_state.n_samples)

    if st.sidebar.button("Generate Ulang Data"):
        update_simulation()

    # ====== STEP 3: HITUNG ISPU ======
    def calculate_ispu_pm25(pm25):
        if pm25 <= 15.5:
            return pm25 * 50 / 15.5
        elif pm25 <= 55.4:
            return 50 + (pm25 - 15.5) * 50 / (55.4 - 15.5)
        elif pm25 <= 150.4:
            return 100 + (pm25 - 55.4) * 100 / (150.4 - 55.4)
        elif pm25 <= 250.4:
            return 200 + (pm25 - 150.4) * 100 / (250.4 - 150.4)
        else:
            return 300 + (pm25 - 250.4) * 100 / (500 - 250.4)

    def calculate_ispu_pm10(pm10):
        if pm10 <= 50:
            return pm10 * 50 / 50
        elif pm10 <= 150:
            return 50 + (pm10 - 50) * 50 / (150 - 50)
        elif pm10 <= 350:
            return 100 + (pm10 - 150) * 100 / (350 - 150)
        elif pm10 <= 420:
            return 200 + (pm10 - 350) * 100 / (420 - 350)
        else:
            return 300 + (pm10 - 420) * 100 / (500 - 420)

    def calculate_ispu_so2(so2):
        if so2 <= 52:
            return so2 * 50 / 52
        elif so2 <= 180:
            return 50 + (so2 - 52) * 50 / (180 - 52)
        elif so2 <= 365:
            return 100 + (so2 - 180) * 100 / (365 - 180)
        elif so2 <= 800:
            return 200 + (so2 - 365) * 100 / (800 - 365)
        else:
            return 300 + (so2 - 800) * 100 / (1200 - 800)

    def calculate_ispu_no2(no2):
        if no2 <= 80:
            return no2 * 50 / 80
        elif no2 <= 200:
            return 50 + (no2 - 80) * 50 / (200 - 80)
        elif no2 <= 1130:
            return 100 + (no2 - 200) * 100 / (1130 - 200)
        elif no2 <= 2000:
            return 200 + (no2 - 1130) * 100 / (2000 - 1130)
        else:
            return 300 + (no2 - 2000) * 100 / (3000 - 2000)

    def calculate_ispu_co(co_ugm3):  # input dalam µg/m³
        if co_ugm3 <= 4000:
            return co_ugm3 * 50 / 4000
        elif co_ugm3 <= 8000:
            return 50 + (co_ugm3 - 4000) * 50 / (8000 - 4000)
        elif co_ugm3 <= 15000:
            return 100 + (co_ugm3 - 8000) * 100 / (15000 - 8000)
        elif co_ugm3 <= 17000:
            return 200 + (co_ugm3 - 15000) * 100 / (17000 - 15000)
        else:
            return 300 + (co_ugm3 - 17000) * 100 / (30000 - 17000)

    def calculate_ispu_o3(o3):
        if o3 <= 120:
            return o3 * 50 / 120
        elif o3 <= 235:
            return 50 + (o3 - 120) * 50 / (235 - 120)
        elif o3 <= 400:
            return 100 + (o3 - 235) * 100 / (400 - 235)
        elif o3 <= 800:
            return 200 + (o3 - 400) * 100 / (800 - 400)
        else:
            return 300 + (o3 - 800) * 100 / (1000 - 800)

    # Hitung semua ISPU
    if 'ISPU_PM2_5' not in st.session_state.df.columns:
        st.session_state.df['ISPU_PM2_5'] = st.session_state.df['PM2.5'].apply(calculate_ispu_pm25).round(2)
        st.session_state.df['ISPU_PM10'] = st.session_state.df['PM10'].apply(calculate_ispu_pm10).round(2)
        st.session_state.df['ISPU_SO2'] = st.session_state.df['SO2'].apply(calculate_ispu_so2).round(2)
        st.session_state.df['ISPU_NO2'] = st.session_state.df['NO2'].apply(calculate_ispu_no2).round(2)
        st.session_state.df['ISPU_CO'] = st.session_state.df['CO'].apply(calculate_ispu_co).round(2)
        st.session_state.df['ISPU_O3'] = st.session_state.df['O3'].apply(calculate_ispu_o3).round(2)

        # Ambil ISPU maksimum sebagai indeks utama
        st.session_state.df['ISPU_max'] = st.session_state.df[['ISPU_PM2_5', 'ISPU_PM10', 'ISPU_SO2', 'ISPU_NO2', 'ISPU_CO', 'ISPU_O3']].max(axis=1)
        st.session_state.df['Parameter_Dominan'] = st.session_state.df[['ISPU_PM2_5', 'ISPU_PM10', 'ISPU_SO2', 'ISPU_NO2', 'ISPU_CO', 'ISPU_O3']].idxmax(axis=1).str[5:]

        st.session_state.df['Polusi_Total'] = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].sum(axis=1)

    # ====== TAMPILAN STREAMLIT ======
    st.sidebar.header("Parameter Simulasi")
    show_raw_data = st.sidebar.checkbox("Tampilkan Data Mentah")

    if show_raw_data:
        st.subheader("Data Mentah Simulasi")
        st.dataframe(st.session_state.df)


    # Tab layout
    tab1, tab2, tab3, tab4 = st.tabs(["Distribusi Parameter", "Visualisasi ISPU", "Kesimpulan", "Parameter Dari File Excel"])

    with tab1:
        st.header("Distribusi Parameter Polutan Simulasi")
        
        # ====== STATISTIK SUMMARY DAN PARAMETER ======
        st.subheader("Ringkasan Statistik Simulasi")
        
        # 1. Statistik Deskriptif (dari data simulasi)
        st.markdown("**Statistik Deskriptif (µg/m³):**")
        summary_stats = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].describe().round(2)
        st.dataframe(summary_stats)
        
        # 2. Parameter Distribusi (dari data Simulasi)
        st.markdown("**Parameter Distribusi (Dihitung dari Data Simulasi):**")
        
        # Buat tabel parameter
        param_table = pd.DataFrame(columns=['Polutan', 'Distribusi', 'Parameter 1', 'Parameter 2'])
        st.session_state.params_simulasi = {}  # <-- Buat variabel baru

        for col in ['PM2.5', 'PM10', 'O3']:
            data = st.session_state.df[col]  # <-- Gunakan data simulasi
            params = weibull_min.fit(data, floc=0)
            st.session_state.params_simulasi[col] = ('Weibull', *params)

        for col in ['SO2', 'NO2', 'CO']:
            mu, std = norm.fit(st.session_state.df[col])
            st.session_state.params_simulasi[col] = ('Normal', mu, std)
        for col in ['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']:
            dist_type, *dist_params = st.session_state.params_simulasi[col]
            if dist_type == 'Weibull':
                param_table.loc[len(param_table)] = [col, 'Weibull', f'Shape k: {dist_params[0]:.2f}', f'Scale λ: {dist_params[2]:.2f}']
            else:
                param_table.loc[len(param_table)] = [col, 'Normal', f'Mean: {dist_params[0]:.2f}', f'Std: {dist_params[1]:.2f}']
        
        st.dataframe(param_table)



        # Plot distribusi parameter
        st.subheader("Distribusi Parameter dengan Kurva Fit")
        fig1, axes1 = plt.subplots(2, 3, figsize=(15, 10))
        fig1.suptitle('Distribusi Parameter Kualitas Udara', fontsize=14)
        axes1 = axes1.ravel()

        st.markdown("""
        **Penjelasan:**
        - Grafik membandingkan distribusi aktual (histogram) dengan kurva teoritis
        - Garis hijau: Distribusi Normal
        - Garis merah: Distribusi Weibull
        - PM2.5, PM10, dan O3 dimodelkan dengan distribusi Weibull
        - SO2, NO2, dan CO dimodelkan dengan distribusi normal
        """)
        
        for i, col in enumerate(['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']):
            dist_type, *dist_params = st.session_state.params[col]
            ax = axes1[i]
            sns.histplot(st.session_state.df[col], bins=20, kde=False, stat="density", alpha=0.6, label='Histogram', ax=ax)
            
            x = np.linspace(0, st.session_state.df[col].max(), 100)
            if dist_type == 'Weibull':
                c, loc, scale = dist_params
                pdf = weibull_min.pdf(x, c, loc, scale)
                ax.plot(x, pdf, 'r-', lw=2, label='Weibull Fit')
            else:
                mu, std = dist_params
                pdf = norm.pdf(x, mu, std)
                ax.plot(x, pdf, 'g-', lw=2, label='Normal Fit')
                
            ax.set_title(f'{col} ({dist_type})')
            ax.set_xlabel('Konsentrasi')
            ax.set_ylabel('Density')
            ax.legend()
        
        plt.tight_layout()
        st.pyplot(fig1)
        st.caption("Histogram menunjukkan distribusi konsentrasi polutan dengan kurva distribusi yang sesuai.")
        plt.close()

    with tab2:
        st.header("Visualisasi Indeks ISPU")
        
        # Tren Harian ISPU
        st.subheader("Tren Harian ISPU")
        fig3 = plt.figure(figsize=(14, 6))

        st.markdown("""
        **Penjelasan Tren ISPU:**
        - Garis biru: Nilai ISPU harian dengan fluktuasi alami
        - Garis merah: Rata-rata 7 hari untuk melihat tren
        - Lonjakan ISPU bisa disebabkan oleh:
        - Peningkatan aktivitas kendaraan
        - Pembakaran biomassa
        - Kondisi meteorologi yang tidak mendispersi polutan
        """)
        
        plt.plot(st.session_state.df.index, st.session_state.df['ISPU_max'], label='ISPU Harian', alpha=0.5)
        plt.plot(st.session_state.df['ISPU_max'].rolling(7).mean(), label='Rata-rata 7 Hari', color='red')
        plt.title('Tren Harian ISPU dengan Smoothing')
        plt.xlabel('Hari')
        plt.ylabel('ISPU')
        plt.grid(True)
        plt.legend()
        st.pyplot(fig3)
        st.caption("Plot menunjukkan tren ISPU harian dengan rata-rata 7 hari terakhir.")
        
        # Kategori ISPU
        st.subheader("Distribusi Kategori ISPU")
        st.markdown("""
        **Kategori ISPU dan Dampaknya:**
        1. **Baik** (Hijau): Aktivitas luar ruangan aman
        2. **Sedang** (Kuning): Kurangi aktivitas panjang di luar bagi penderita pernapasan
        3. **Tidak Sehat** (Oranye): Hindari aktivitas luar ruangan berkepanjangan
        4. **Sangat Tidak Sehat** (Merah): Hindari semua aktivitas luar ruangan
        5. **Berbahaya** (Merah Tua): Tetap di dalam ruangan dengan pembersih udara
        """)
        # Hitung jumlah hari per kategori
        kategori_urut = ['Baik', 'Sedang', 'Tidak Sehat', 'Sangat Tidak Sehat', 'Berbahaya']
        st.session_state.df['Kategori_ISPU'] = st.session_state.df['ISPU_max'].apply(lambda x: "Baik" if x <= 50 else 
                                                "Sedang" if x <= 100 else 
                                                "Tidak Sehat" if x <= 200 else 
                                                "Sangat Tidak Sehat" if x <= 300 else 
                                                "Berbahaya")
        kategori_counts = st.session_state.df['Kategori_ISPU'].value_counts().reindex(kategori_urut, fill_value=0)
        
        # Warna dan style
        warna_kategori = {
            'Baik': 'green',
            'Sedang': 'gold',
            'Tidak Sehat': 'orange',
            'Sangat Tidak Sehat': 'red',
            'Berbahaya': 'darkred'
        }
        colors = [warna_kategori[k] for k in kategori_counts.index]
        
        # Filter out categories with 0 count
        filtered_counts = kategori_counts[kategori_counts > 0]
        filtered_colors = [colors[i] for i in range(len(kategori_counts)) if kategori_counts[i] > 0]

        # Create visualizations only if there are categories to show
        if not filtered_counts.empty:
            fig5, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))
            
            # Pie Chart
            ax1.pie(filtered_counts, labels=filtered_counts.index,
                    autopct='%1.1f%%', startangle=90, colors=filtered_colors,
                    shadow=True, textprops={'color': 'black', 'fontsize': 10})
            ax1.set_title('Distribusi Persentase Kategori ISPU (Tanpa Kategori 0%)')
            ax1.axis('equal')
            
            # Bar Chart
            bars = ax2.bar(filtered_counts.index, filtered_counts.values, color=filtered_colors, edgecolor='black')
            ax2.set_title('Jumlah Hari per Kategori ISPU (Tanpa Kategori 0%)')
            ax2.set_xlabel('Kategori ISPU')
            ax2.set_ylabel('Jumlah Hari')
            ax2.grid(axis='y', linestyle='--', alpha=0.7)
            ax2.set_axisbelow(True)
            for bar in bars:
                height = bar.get_height()
                ax2.text(bar.get_x() + bar.get_width()/2., height + 0.5, 
                        f'{int(height)}\n({height/len(st.session_state.df):.1%})',
                        ha='center', va='bottom', fontsize=9)
            
            plt.tight_layout()
            st.pyplot(fig5)
            st.caption("Diagram lingkaran dan batang menunjukkan distribusi kategori ISPU selama periode simulasi (kategori dengan 0% tidak ditampilkan).")
        else:
            st.warning("Tidak ada data kategori ISPU yang tersedia untuk ditampilkan.")
        
        # ====== 4 DIAGRAM BARU ======
        
        # 2. Bar Chart Polutan Dominan (menggunakan variabel utama)
        st.subheader("Polutan Dominan Penyebab ISPU Tertinggi")
        st.markdown("""
        **Penjelasan Diagram:**
        - **Sumbu X**: Nama polutan (PM2.5, PM10, SO2, NO2, CO, O3)
        - **Sumbu Y**: Jumlah hari polutan tersebut menjadi penyebab ISPU tertinggi
        - Warna biru muda: Frekuensi dominansi masing-masing polutan
        - Garis grid: Membantu membaca nilai secara akurat
        """)
        fig_bar = plt.figure(figsize=(10, 6))
        polutan_dominan = st.session_state.df['Parameter_Dominan'].value_counts()
        polutan_dominan.plot(kind='bar', color='skyblue')
        plt.title('Frekuensi Polutan Dominan Penyebab ISPU Tertinggi')
        plt.xlabel('Polutan')
        plt.ylabel('Jumlah Hari')
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        st.pyplot(fig_bar)

        
        # 3. Line Chart Tren Polutan Utama (diagram bebas)
        st.subheader("Tren Polutan Utama")
        st.markdown("""
        **Penjelasan Diagram:**
        - **Sumbu X**: Hari ke- (dikonversi dari hari simulasi)
        - **Sumbu Y**: Konsentrasi harian (µg/m³)
        - Garis berwarna: Mewakili polutan berbeda
        - Area terarsir: Variasi harian dalam hari tersebut
        - Pola mingguan membantu identifikasi pengaruh aktivitas manusia
        """)
        fig_line = plt.figure(figsize=(14, 6))
        for polutan in ['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']:
            plt.plot(st.session_state.df.index, st.session_state.df[polutan], label=polutan, alpha=0.7)
        plt.title('Tren Polutan Utama')
        plt.xlabel('Hari')
        plt.ylabel('Konsentrasi')
        plt.legend()
        plt.grid(True)
        st.pyplot(fig_line)

        
        # 4. Histogram PM2.5 (diagram bebas)
        st.subheader("Distribusi PM2.5")
        st.markdown("""
        **Penjelasan Diagram:**
        - **Sumbu X**: Rentang konsentrasi PM2.5 (µg/m³)
        - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
        - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
        - Garis KDE (Kernel Density Estimate): Estimasi kurva distribusi
        - Garis Merah batas aman ISPU
        """)
        fig_hist = plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['PM2.5'], bins=30, kde=True, color='teal')
        plt.axvline(x=15.5, color='red', linestyle='--', label='Batas Aman ISPU')    
        plt.title('Distribusi Konsentrasi PM2.5')
        plt.xlabel('PM2.5 (µg/m³)')
        plt.ylabel('Frekuensi')
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        st.pyplot(fig_hist)

        # 5. Histogram PM10
        st.subheader("Distribusi PM10")
        st.markdown("""
        **Penjelasan Diagram:**
        - **Sumbu X**: Rentang konsentrasi PM10 (µg/m³)
        - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
        - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
        - Garis KDE (Kernel Density Estimate): Estimasi kurva distribusi
        - Garis Merah batas aman ISPU
        """)
        fig_hist = plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['PM10'], bins=30, kde=True, color='teal')
        plt.axvline(x=50, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi PM10')
        plt.xlabel('PM10 (µg/m³)')
        plt.ylabel('Frekuensi')
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        st.pyplot(fig_hist)

        # 6. Histogram SO2
        st.subheader("Distribusi SO2")
        st.markdown("""
        **Penjelasan Diagram:**
        - **Sumbu X**: Rentang konsentrasi SO2 (µg/m³)
        - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
        - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
        - Garis KDE: Estimasi kurva distribusi
        - Garis Merah batas aman ISPU
        """)
        fig_hist = plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['SO2'], bins=30, kde=True, color='teal')
        plt.axvline(x=35, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi SO2')
        plt.xlabel('SO2 (µg/m³)')
        plt.ylabel('Frekuensi')
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        st.pyplot(fig_hist)

        # 7. Histogram NO2
        st.subheader("Distribusi NO2")
        st.markdown("""
        **Penjelasan Diagram:**
        - **Sumbu X**: Rentang konsentrasi NO2 (µg/m³)
        - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
        - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
        - Garis KDE: Estimasi kurva distribusi
        - Garis Merah batas aman ISPU
        """)
        fig_hist = plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['NO2'], bins=30, kde=True, color='teal')
        plt.axvline(x=80, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi NO2')
        plt.xlabel('NO2 (µg/m³)')
        plt.ylabel('Frekuensi')
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        st.pyplot(fig_hist)

        # 8. Histogram CO
        st.subheader("Distribusi CO")
        st.markdown("""
        **Penjelasan Diagram:**
        - **Sumbu X**: Rentang konsentrasi CO (µg/m³)
        - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
        - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
        - Garis KDE: Estimasi kurva distribusi
        - Garis Merah batas aman ISPU
        """)
        fig_hist = plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['CO'], bins=30, kde=True, color='teal')
        plt.axvline(x=4, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi CO')
        plt.xlabel('CO (mg/m³)')
        plt.ylabel('Frekuensi')
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        st.pyplot(fig_hist)

        # 9. Histogram O3
        st.subheader("Distribusi O3")
        st.markdown("""
        **Penjelasan Diagram:**
        - **Sumbu X**: Rentang konsentrasi O3 (µg/m³)
        - **Sumbu Y**: Frekuensi kemunculan (jumlah hari)
        - Batang histogram: Menunjukkan seberapa sering suatu konsentrasi terjadi
        - Garis KDE: Estimasi kurva distribusi
        - Garis Merah batas aman ISPU
        """)
        fig_hist = plt.figure(figsize=(10, 6))
        sns.histplot(st.session_state.df['O3'], bins=30, kde=True, color='teal')
        plt.axvline(x=120, color='red', linestyle='--', label='Batas Aman ISPU')
        plt.title('Distribusi Konsentrasi O3')
        plt.xlabel('O3 (µg/m³)')
        plt.ylabel('Frekuensi')
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        st.pyplot(fig_hist)

    with tab3:
        st.header("Kesimpulan Simulasi")
        
        # Hitung statistik penting
        kategori_ispu_counts = st.session_state.df['Kategori_ISPU'].value_counts()
        rata2_ispu = st.session_state.df['ISPU_max'].mean()
        max_ispu = st.session_state.df['ISPU_max'].max()
        hari_tidak_sehat = len(st.session_state.df[st.session_state.df['Kategori_ISPU'].isin(['Tidak Sehat', 'Sangat Tidak Sehat', 'Berbahaya'])])
        polusi_tertinggi = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].max().idxmax()
        hari_dengan_polusi_tinggi = len(st.session_state.df[st.session_state.df['Polusi_Total'] > st.session_state.df['Polusi_Total'].quantile(0.75)])


        pm25_over_limit = len(st.session_state.df[st.session_state.df['PM2.5'] > 15.5])
        pm10_over_limit = len(st.session_state.df[st.session_state.df['PM10'] > 50])
        o3_over_limit = len(st.session_state.df[st.session_state.df['O3'] > 120])
        co_over_limit = len(st.session_state.df[st.session_state.df['CO'] > 5])
        so2_over_limit = len(st.session_state.df[st.session_state.df['SO2'] > 20])
        no2_over_limit = len(st.session_state.df[st.session_state.df['NO2'] > 30])
        max_pollutant = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].max().idxmax()
        min_pollutant = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].min().idxmax()
        kategori_urut = ['Baik', 'Sedang', 'Tidak Sehat', 'Sangat Tidak Sehat', 'Berbahaya']
        kategori_counts = st.session_state.df['Kategori_ISPU'].value_counts().reindex(kategori_urut, fill_value=0)
        mean_pollutants = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].mean().sort_values()
        least_harmful = mean_pollutants.idxmin()
        min_pollutants = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].min().sort_values()
        cleanest_pollutant = min_pollutants.idxmin()
        std_pollutants = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].std().sort_values()
        most_stable = std_pollutants.idxmin()


        # Tampilkan kesimpulan
        st.write(f"""
        1. Rata-rata nilai ISPU selama {st.session_state.n_samples} hari adalah {rata2_ispu:.2f}, menunjukkan bahwa kualitas udara cenderung berada pada kategori {'Baik/Sedang' if rata2_ispu < 100 else 'Tidak Sehat'}.
        2. Nilai ISPU tertinggi adalah {max_ispu:.2f}, masuk dalam kategori '{st.session_state.df.loc[st.session_state.df['ISPU_max'].idxmax(), 'Kategori_ISPU']}'. Ini merupakan kondisi kritis yang memerlukan antisipasi lebih lanjut.
        3. Terdapat {hari_tidak_sehat} hari dengan kategori ISPU tidak sehat/sangat tidak sehat/berbahaya. Kelompok rentan seperti lansia dan penderita penyakit paru-paru perlu waspada.
        4. Polutan dominan berdasarkan konsentrasi tertinggi adalah '{polusi_tertinggi}', menunjukkan perlunya fokus mitigasi pada parameter tersebut.
        5. Terdapat {hari_dengan_polusi_tinggi} hari dengan polusi total sangat tinggi (>75% kuartil), yang mungkin dipengaruhi oleh peningkatan beberapa polutan secara bersamaan.
        6. PM2.5 melebihi batas aman ISPU (15.5 µg/m³) sebanyak {pm25_over_limit} hari, menunjukkan perlunya pengendalian emisi kendaraan dan industri.
        7. PM10 melebihi batas aman ISPU (50 µg/m³) sebanyak {pm10_over_limit} hari, menunjukkan polusi udara yang cukup signifikan di wilayah simulasi.
        8. Ozon (O3) melebihi ambang batas aman sebanyak {o3_over_limit} hari.
        9. CO melebihi ambang batas sebanyak {co_over_limit} hari.
        10. SO2 melebihi ambang batas sebanyak {so2_over_limit} hari.
        11. NO2 melebihi ambang batas sebanyak {no2_over_limit} hari.
        12. Polutan {max_pollutant} memiliki konsentrasi tertinggi di seluruh simulasi, menunjukkan bahwa polutan tersebut paling berkontribusi terhadap penurunan kualitas udara.
        13. Polutan {min_pollutant} memiliki konsentrasi tertinggi di minimum.
        14. {least_harmful} merupakan polutan dengan konsentrasi rata-rata terendah (**{mean_pollutants[least_harmful]:.2f} µg/m³**). Sumber emisi {least_harmful} relatif lebih terkendali.  
        15. {cleanest_pollutant} mencapai level terendah (**{min_pollutants[cleanest_pollutant]:.2f} µg/m³**) saat udara paling bersih.
        16. {most_stable} memiliki fluktuasi harian paling stabil (deviasi **{std_pollutants[most_stable]:.2f} µg/m³**).  
        """)
        
        st.write("""
        **Rekomendasi:**
        - Pengurangan emisi kendaraan bermotor dan aktivitas industri
        - Edukasi publik tentang cara melindungi diri dari paparan polusi udara
        - Peningkatan pemantauan kualitas udara di area dengan risiko tinggi
        - Fokus pada pengendalian polutan dominan yang teridentifikasi
        - Implementasi sistem peringatan dini ketika konsentrasi polutan mencapai level berbahaya
        - Fokus pada pengendalian polutan dengan **rata-rata tertinggi** terlebih dahulu  
        - Manfaatkan data polutan **paling stabil** sebagai baseline pemantauan  
        - Gunakan kriteria **minimum terendah** untuk menetapkan target kualitas udara bersih 
        """)

    with tab4:
        st.header("Distribusi Parameter Polutan Excel")
        
        # ====== STATISTIK SUMMARY DAN PARAMETER ======
        st.subheader("Ringkasan Statistik Excel")
        
        # 1. Statistik Deskriptif (dari data Excel)
        st.markdown("**Statistik Deskriptif (µg/m³):**")
        summary_stats = st.session_state.df_excel[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].describe().round(2)
        st.dataframe(summary_stats)
        
        # 2. Parameter Distribusi (dari data Excel)
        st.markdown("**Parameter Distribusi (Dihitung dari Data Excel):**")
        
        # Buat tabel parameter
        param_table = pd.DataFrame(columns=['Polutan', 'Distribusi', 'Parameter 1', 'Parameter 2'])
        
        for col in ['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']:
            dist_type, *dist_params = st.session_state.params[col]
            if dist_type == 'Weibull':
                param_table.loc[len(param_table)] = [col, 'Weibull', f'Shape k: {dist_params[0]:.2f}', f'Scale λ: {dist_params[2]:.2f}']
            else:
                param_table.loc[len(param_table)] = [col, 'Normal', f'Mean: {dist_params[0]:.2f}', f'Std: {dist_params[1]:.2f}']
        
        st.dataframe(param_table)
        

    # Download buttons
    st.sidebar.header("Unduh Hasil")
    if st.sidebar.button("Buat Laporan Word"):
        with st.spinner('Membuat dokumen Word...'):
            doc_buffer = create_word_document(st.session_state.df, st.session_state.params, st.session_state.n_samples)
            st.sidebar.download_button(
                label="Unduh Laporan Word",
                data=doc_buffer,
                file_name=f"Laporan_ISPU_{st.session_state.n_samples}_hari.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    @st.cache_data
    def convert_df(df):
        return df.to_csv(index=False).encode('utf-8')

    csv = convert_df(st.session_state.df)

    st.sidebar.download_button(
        label="Unduh Data Simulasi (CSV)",
        data=csv,
        file_name=f"simulasi_ispu_{st.session_state.n_samples}_hari.csv",
        mime="text/csv"
    )
    pass

elif page == "Analisis Data":
    st.title("Analisis Data Lanjutan")
    
    # Ensure Kategori_ISPU column exists
    if 'Kategori_ISPU' not in st.session_state.df.columns:
        st.session_state.df['Kategori_ISPU'] = st.session_state.df['ISPU_max'].apply(
            lambda x: "Baik" if x <= 50 else 
                     "Sedang" if x <= 100 else 
                     "Tidak Sehat" if x <= 200 else 
                     "Sangat Tidak Sehat" if x <= 300 else 
                     "Berbahaya"
        )
    
    tab1, tab3, tab4 = st.tabs(["Filter Data", "Korelasi Polutan", "Analisis Temporal"])
    
    with tab1:
        st.header("Filter Data Simulasi")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Filter by ISPU category
            st.subheader("Filter Berdasarkan Kategori ISPU")
            categories = st.multiselect(
                "Pilih Kategori ISPU:",
                options=["Baik", "Sedang", "Tidak Sehat", "Sangat Tidak Sehat", "Berbahaya"],
                default=["Baik", "Sedang", "Tidak Sehat"]
            )
            
        
        with col2:
            # Filter by date range
            st.subheader("Filter Berdasarkan Rentang Hari")
            min_day = st.number_input("Hari Awal", min_value=0, max_value=st.session_state.n_samples, value=0)
            max_day = st.number_input("Hari Akhir", min_value=0, max_value=st.session_state.n_samples, value=st.session_state.n_samples)
            
            # Filter by ISPU value
            st.subheader("Filter Berdasarkan Nilai ISPU")
            min_ispu, max_ispu = st.slider(
                "Rentang Nilai ISPU",
                min_value=0,
                max_value=int(st.session_state.df['ISPU_max'].max()) + 50,
                value=(0, int(st.session_state.df['ISPU_max'].max()))
            )
        
        # Apply filters
        filtered_df = st.session_state.df.copy()
        
        if categories:
            filtered_df = filtered_df[filtered_df['Kategori_ISPU'].isin(categories)]

        
        filtered_df = filtered_df[(filtered_df.index >= min_day) & (filtered_df.index <= max_day)]
        filtered_df = filtered_df[(filtered_df['ISPU_max'] >= min_ispu) & (filtered_df['ISPU_max'] <= max_ispu)]
        
        st.subheader("Data Terfilter")
        st.dataframe(filtered_df)
        
        # Show statistics for filtered data
        st.subheader("Statistik untuk Data Terfilter")
        st.write(filtered_df.describe())
        
        # Export filtered data
        st.download_button(
            label="Unduh Data Terfilter (CSV)",
            data=filtered_df.to_csv(index=False).encode('utf-8'),
            file_name="data_ispu_terfilter.csv",
            mime="text/csv"
        )
    
    with tab3:
        st.header("Analisis Korelasi Antar Polutan")
        
        # Correlation analysis
        st.subheader("Matriks Korelasi")
        corr_matrix = st.session_state.df[['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3']].corr()
        
        # Heatmap visualization
        fig, ax = plt.subplots(figsize=(10, 8))
        sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', center=0, ax=ax)
        st.pyplot(fig)
        
        # Interpretation
        st.markdown("""
        **Interpretasi Korelasi:**
        - Nilai mendekati +1: Korelasi positif kuat
        - Nilai mendekati -1: Korelasi negatif kuat
        - Nilai mendekati 0: Tidak ada korelasi
        """)
    
    with tab4:
        st.header("Analisis Temporal")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Analisis Musiman")
            selected_pollutant = st.selectbox(
                "Pilih Polutan:",
                options=['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3', 'ISPU_max'],
                index=0,
                key="temp_poll"
            )
            
            resample_freq = st.selectbox(
                "Frekuensi Resampling:",
                options=["Harian", "Mingguan", "Bulanan"],
                index=1
            )
            
            # Convert resample frequency
            freq_map = {
                "Harian": "D",
                "Mingguan": "W",
                "Bulanan": "M"
            }
            
            # Create temporal analysis
            temp_df = filtered_df.copy()
            temp_df['Date'] = pd.to_datetime('2022-12-01') + pd.to_timedelta(temp_df.index, unit='D')
            temp_df.set_index('Date', inplace=True)
            
            resampled = temp_df.resample(freq_map[resample_freq])[selected_pollutant].mean()
            
        with col2:
            st.subheader("Visualisasi Tren")
            plot_type = st.selectbox(
                "Jenis Plot Temporal:",
                options=["Line Plot", "Area Plot", "Bar Plot"],
                index=0
            )
            
            fig, ax = plt.subplots(figsize=(10, 6))
            
            if len(resampled) > 0:  # Only plot if we have data
                if plot_type == "Line Plot":
                    resampled.plot(ax=ax, title=f'Tren {selected_pollutant} ({resample_freq})', marker='o')
                elif plot_type == "Area Plot":
                    resampled.plot.area(ax=ax, title=f'Tren {selected_pollutant} ({resample_freq})', alpha=0.4)
                elif plot_type == "Bar Plot":
                    resampled.plot.bar(ax=ax, title=f'Tren {selected_pollutant} ({resample_freq})')
                
                plt.xticks(rotation=45)
                plt.grid(True)
                st.pyplot(fig)
            else:
                st.warning("Tidak ada data yang tersedia untuk ditampilkan setelah filtering.")
                st.pyplot(fig)  # Show empty plot
            
        st.subheader("Analisis Periodik")
        st.markdown("""
        **Pola yang Mungkin Teramati:**
        - **Pola Harian**: Fluktuasi konsentrasi polutan sepanjang hari
        - **Pola Mingguan**: Perbedaan antara hari kerja dan akhir pekan
        - **Pola Musiman**: Variasi konsentrasi antar musim
        
        **Faktor yang Mempengaruhi:**
        - Aktivitas manusia (jam sibuk, hari kerja vs akhir pekan)
        - Kondisi meteorologi (inversi suhu, kecepatan angin)
        - Emisi musiman (pemanas ruangan, pembakaran biomassa)
        """)

elif page == "Visualisasi Lanjutan":
    st.title("Visualisasi Lanjutan Kualitas Udara")
    
    tab1, tab3 = st.tabs(["Time Series", "Scatter Plot"])
    
    with tab1:
        st.header("Analisis Time Series")
        
        # Select pollutants to plot
        selected_pollutants = st.multiselect(
            "Pilih Polutan untuk Ditampilkan:",
            options=['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3'],
            default=['PM2.5', 'PM10'],
            key="ts_select"
        )
        
        # Rolling average option
        rolling_window = st.slider("Window Rata-rata Bergerak (hari):", 1, 30, 7, key="roll_window")
        
        if selected_pollutants:
            # Create the plot
            fig, ax = plt.subplots(figsize=(12, 6))
            for pollutant in selected_pollutants:
                ax.plot(st.session_state.df.index, 
                       st.session_state.df[pollutant].rolling(rolling_window).mean(),
                       label=pollutant)
            
            ax.set_title(f"Tren Polutan (Rata-rata {rolling_window} Hari)")
            ax.set_xlabel("Hari")
            ax.set_ylabel("Konsentrasi (µg/m³)")
            ax.legend()
            ax.grid(True)
            st.pyplot(fig)
    
    with tab3:
        st.header("Scatter Plot Hubungan Polutan")
        
        # Select x and y axes
        col1, col2 = st.columns(2)
        with col1:
            x_axis = st.selectbox(
                "Sumbu X:",
                options=['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3'],
                index=0,
                key="scatter_x"
            )
        with col2:
            y_axis = st.selectbox(
                "Sumbu Y:",
                options=['PM2.5', 'PM10', 'SO2', 'NO2', 'CO', 'O3'],
                index=1,
                key="scatter_y"
            )
        
        # Create scatter plot
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.scatterplot(
            data=st.session_state.df,
            x=x_axis,
            y=y_axis,
            hue='Kategori_ISPU',
            palette={
                'Baik': 'green',
                'Sedang': 'yellow',
                'Tidak Sehat': 'orange',
                'Sangat Tidak Sehat': 'red',
                'Berbahaya': 'purple'
            },
            ax=ax
        )
        ax.set_title(f"Hubungan antara {x_axis} dan {y_axis}")
        st.pyplot(fig)

